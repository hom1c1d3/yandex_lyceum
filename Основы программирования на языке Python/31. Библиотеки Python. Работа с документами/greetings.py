from docx import Document
from sys import stdin
from docx.shared import RGBColor

document = Document()
place, date, *names = stdin.read().split('\n')

for name in names:
    document.add_heading('Приглашение', 0)
    p = document.add_paragraph('Здравствуй, ')
    run = p.add_run(name + '!')
    run.font.name = 'Comic Sans MS'
    run.bold = True
    p.add_run(' В этот понедельник, ')
    run = p.add_run('8 марта ')
    run.font.color.rgb = RGBColor(255, 0, 0)
    run = p.add_run(place)
    run.font.name = 'Times New Roman'
    run.italic = True
    p.add_run(f' состоится концерт, посвященный международному женскому дню, ждем тебя {date.split()[0]} ')
    run = p.add_run(date.split()[1] + '.')
    run.font.name = 'Papyrus'
    run.underline = True
    document.add_page_break()

document.save('greet.docx')