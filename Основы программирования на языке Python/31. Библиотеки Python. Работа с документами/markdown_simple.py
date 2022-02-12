from itertools import groupby

from docx import Document


def markdown_to_docx(text: str):
    document = Document()

    text = text.splitlines()
    document.add_heading(text[0], level=0)
    text = text[1:]
    for line in text:
        if not line:
            document.add_paragraph()
            continue
        comm = [''.join(v) for k, v in groupby(line, key=lambda x: x == '#')][0]
        if '#' in comm:
            document.add_heading(line[len(comm) + 1:], level=comm.count('#'))

        elif len(line) > 2 and (line[0].isdigit() and line[1] == '.'):
            document.add_paragraph(line[3:], style='List Number')
        elif line[:2] == line[-2:]:
            char = line[0]
            num_char = line.count(char) // 2
            line = line.strip(char)
            run = document.add_paragraph().add_run(line)
            run.bold = num_char > 1
            num_char -= 2
            run.italic = num_char == 1
        elif line[0] == line[-1]:
            document.add_paragraph().add_run(line[1:-1]).italic = True
        elif line[0] in ('-', '*', '+'):
            document.add_paragraph(line[2:], style='List Bullet')
        else:
            document.add_paragraph(line)
    document.save('res.docx')


markdown_to_docx('''test02
Абзацы создаются при помощи пустой строки. Если вокруг текста сверху и снизу есть пустые строки, то текст превращается в абзац.

Чтобы сделать перенос строки вместо абзаца,
нужно поставить два пробела в конце предыдущей строки.

Заголовки отмечаются диезом `#` в начале строки, от одного до шести. Например:

# Заголовок первого уровня
## Заголовок h2
### Заголовок h3
#### Заголовок h4
##### Заголовок h5
###### Заголовок h6

В декоративных целях заголовки можно «закрывать» с обратной стороны.

### Списки

Для разметки неупорядоченных списков можно использовать или `*`, или `-`, или `+`:

- элемент 1
- элемент 2
- элемент ...

Упорядоченный список:

1. элемент 1
2. элемент 2
3. элемент 3
4. QT 1. Что такое QT и PyQT. Знакомство. Donec sit amet nisl. Aliquam semper ipsum sit amet velit. Suspendisse id sem consectetuer libero luctus adipiscing.

_italic_
и это тоже
*italic*
А вот так уже
__strong__
, и так тоже
**strong**
***strong_italic***''')